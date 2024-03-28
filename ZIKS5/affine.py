import os
import docx
import time
import random

UPPERLETTERS = 'ABCDEFGHIJKLMNOPQRSTUVWXYZ'
LETTERS_AND_SPACE = UPPERLETTERS + UPPERLETTERS.lower() + ' \t\n'
SILENT_MODE = True
SYMBOLS = 'ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz'


def gcd(a, b):
    while a != 0:
        a, b = b % a, a
    return b


def find_mod_inverse(a, m):
    if gcd(a, m) != 1:
        return None

    u1, u2, u3 = 1, 0, a
    v1, v2, v3 = 0, 1, m
    while v3 != 0:
        q = u3 // v3
        v1, v2, v3, u1, u2, u3 = (u1 - q * v1), (u2 - q * v2), (u3 - q * v3), v1, v2, v3
    return u1 % m


def get_key_parts(key):
    key_a = key // len(SYMBOLS)
    key_b = key % len(SYMBOLS)
    return key_a, key_b


def load_dictionary():
    dictionary_file = open('dictionary.txt')
    english_words = {}
    for word in dictionary_file.read().split('\n'):
        english_words[word] = None
    dictionary_file.close()
    return english_words


ENGLISH_WORDS = load_dictionary()


def get_english_count(message):
    message = message.upper()
    message = remove_non_letters(message)
    possible_words = message.split()

    if not possible_words:
        return 0.0

    matches = 0
    for word in possible_words:
        if word in ENGLISH_WORDS:
            matches += 1
    return float(matches) / len(possible_words)


def remove_non_letters(message):
    letters_only = []
    for symbol in message:
        if symbol in LETTERS_AND_SPACE:
            letters_only.append(symbol)
    return ''.join(letters_only)


def is_english(message, word_percentage=20, letter_percentage=85):
    words_match = get_english_count(message) * 100 >= word_percentage
    num_letters = len(remove_non_letters(message))
    message_letters_percentage = float(num_letters) / len(message) * 100
    letters_match = message_letters_percentage >= letter_percentage
    return words_match and letters_match


def main():
    global translated
    source_choice = input("Choose the source:\n1. Console\n2. File (.docx)\nEnter your choice: ")
    if source_choice == '1':
        text = input("Enter the text: ")
    elif source_choice == '2':
        file_path = input("Enter the file path: ")
        if not file_path.endswith('.docx'):
            file_path += '.docx'
        if not os.path.exists(file_path):
            print("File not found.")
            return
        text = read_docx_file(file_path)
    else:
        print("Invalid choice.")
        return

    choice = input(
        "Choose an option:\n1. Encode\n2. Decode with known key\n3. Hack message and key\nEnter your choice: ")

    if choice == '1':
        key_choice = input("Choose an option:\n1. Enter key manually\n2. Generate random key\nEnter your choice: ")
        if key_choice == '1':
            the_key = int(input("Enter the key (an integer): "))
        elif key_choice == '2':
            the_key = get_random_key()
        else:
            print("Invalid choice.")
            return

        choice = input(
            "Choose an option:\n1. Single affine cipher\n2. Double affine cipher\nEnter your choice: ")
        if choice == '1':
            start_time = time.perf_counter()
            translated = encrypt_message(the_key, text)
            end_time = time.perf_counter()
            duration = end_time - start_time
            print(f"Encryption time: {duration} seconds")
        elif choice == '2':
            start_time = time.perf_counter()
            translated = double_encrypt_message(the_key, text)
            end_time = time.perf_counter()
            duration = end_time - start_time
            print(f"Encryption time: {duration} seconds")
        else:
            print("Invalid choice.")
            return
        print('Key: %s' % the_key)
        print("Encoded message: ", translated)
        save_choice = input("Do you want to save the encoded message to a .docx file? (yes/no): ")
        if save_choice.lower() == 'yes':
            save_to_docx(translated)
    elif choice == '2':
        key_choice = input("Choose an option:\n1. Enter key manually\n2. Generate random key\nEnter your choice: ")
        if key_choice == '1':
            the_key = int(input("Enter the key (an integer): "))
        elif key_choice == '2':
            the_key = get_random_key()
        else:
            print("Invalid choice.")
            return

        choice = input(
            "Choose an option:\n1. Single affine cipher\n2. Double affine cipher\nEnter your choice: ")
        if choice == '1':
            start_time = time.perf_counter()
            translated = decrypt_message(the_key, text)
            end_time = time.perf_counter()
            duration = end_time - start_time
            print(f"Decryption time: {duration} seconds")
        elif choice == '2':
            start_time = time.perf_counter()
            translated = double_decrypt_message(the_key, text)
            end_time = time.perf_counter()
            duration = end_time - start_time
            print(f"Decryption time: {duration} seconds")
        else:
            print("Invalid choice.")
            return
        print('Key: %s' % the_key)
        print("Decoded message: ", translated)
    elif choice == '3':
        hacked_message = hack_affine(text)
        if hacked_message is not None:
            print(hacked_message)
        else:
            print('Failed to hack encryption.')
    else:
        print("Invalid choice!")


def get_random_key():
    while True:
        key_a = random.randint(2, len(SYMBOLS))
        key_b = random.randint(2, len(SYMBOLS))
        if gcd(key_a, len(SYMBOLS)) == 1:
            return key_a * len(SYMBOLS) + key_b


def check_keys(key_a, key_b, mode):
    if key_a == 1 and mode == 'encrypt':
        sys.exit('The affine cipher becomes incredibly weak when key A is set to 1. Choose a different key.')
    if key_b == 0 and mode == 'encrypt':
        sys.exit('The affine cipher becomes incredibly weak when key B is set to 0. Choose a different key.')
    if key_a < 0 or key_b < 0 or key_b > len(SYMBOLS) - 1:
        sys.exit('Key A must be greater than 0 and Key B must be between 0 and %s.' % (len(SYMBOLS) - 1))
    if gcd(key_a, len(SYMBOLS)) != 1:
        sys.exit('Key A (%s) and the symbol set size (%s) are not relatively prime. Choose a different key.' % (key_a, len(SYMBOLS)))
    return


def double_encrypt_message(the_key, plaintext):
    first_encryption = encrypt_message(the_key, plaintext)
    return encrypt_message(the_key, first_encryption)


def encrypt_message(the_key, plaintext):
    key_a, key_b = get_key_parts(the_key)
    check_keys(key_a, key_b, 'encrypt')
    ciphertext = ''
    for letters in plaintext:
        if letters in SYMBOLS:
            letter_index = SYMBOLS.find(letters)
            ciphertext += SYMBOLS[(letter_index * key_a + key_b) % len(SYMBOLS)]
        else:
            ciphertext += letters
    return ciphertext


def double_decrypt_message(the_key, ciphertext):
    first_decryption = decrypt_message(the_key, ciphertext)
    return decrypt_message(the_key, first_decryption)


def decrypt_message(the_key, ciphertext):
    key_a, key_b = get_key_parts(the_key)
    check_keys(key_a, key_b, 'decrypt')
    plaintext = ''
    mod_inverse_of_key_a = find_mod_inverse(key_a, len(SYMBOLS))

    for letters in ciphertext:
        if letters in SYMBOLS:
            letter_index = SYMBOLS.find(letters)
            plaintext += SYMBOLS[(letter_index - key_b) * mod_inverse_of_key_a % len(SYMBOLS)]
        else:
            plaintext += letters
    return plaintext


def hack_affine(ciphertext):
    start_time = time.perf_counter()
    print('Hacking...')
    for key in range(len(SYMBOLS) ** 2):
        key_a = get_key_parts(key)[0]
        if gcd(key_a, len(SYMBOLS)) != 1:
            continue

        decrypted_text = decrypt_message(key, ciphertext)

        if is_english(decrypted_text):
            end_time = time.perf_counter()
            duration = end_time - start_time
            print(f"Hack time: {duration} seconds")
            print()
            print('Possible encryption hack: ')
            print('Key: %s' % key)
            print('Decrypted message: ' + decrypted_text[:300])
            print()
            print('Enter D for done, or just press Enter to continue hacking:')
            response = input('> ')

            if response.strip().upper().startswith('D'):
                return decrypted_text
    return None


def read_docx_file(file_path):
    doc = docx.Document(file_path)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return '\n'.join(full_text)


def save_to_docx(text):
    file_name = input("Enter the file name to save the encoded message: ")
    doc = docx.Document()
    doc.add_paragraph(text)
    doc.save(f"{file_name}.docx")


if __name__ == '__main__':
    main()
