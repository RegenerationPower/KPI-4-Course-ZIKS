import codecs
import os
import random
import string
from docx import Document
import time


def generate_random_key(length):
    return ''.join(random.choice(string.ascii_letters + string.digits) for _ in range(length))


def ksa(key):
    key_length = len(key)
    s = list(range(256))
    j = 0
    for i in range(256):
        j = (j + s[i] + key[i % key_length]) % 256
        s[i], s[j] = s[j], s[i]
    return s


def prga(s):
    i = 0
    j = 0
    while True:
        i = (i + 1) % 256
        j = (j + s[i]) % 256
        s[i], s[j] = s[j], s[i]
        k = s[(s[i] + s[j]) % 256]
        yield k


def get_keystream(key):
    s = ksa(key)
    return prga(s)


def encrypt_logic(key, text_bytes):
    key_bytes = key.encode('utf-8')
    keystream = get_keystream(key_bytes)
    res = []
    for byte in text_bytes:
        val = ("%02X" % (byte ^ next(keystream)))
        res.append(val)
    return ''.join(res)


def encrypt(key, plaintext):
    plaintext_bytes = plaintext.encode('utf-8')
    return encrypt_logic(key, plaintext_bytes)


def decrypt_logic(key, ciphertext):
    res = encrypt_logic(key, codecs.decode(ciphertext, 'hex_codec'))
    characters = []

    for i in range(0, len(res), 2):
        characters.append(chr(int(res[i:i + 2], 16)))

    plaintext = ''.join(characters)
    plaintext = plaintext.encode('latin1').decode('utf-8')
    return plaintext


def decrypt(key, ciphertext):
    return decrypt_logic(key, ciphertext)


def encrypt_file(key, input_file, output_file):
    if input_file.endswith('.docx'):
        doc = Document(input_file)
        plaintext = '\n'.join(paragraph.text for paragraph in doc.paragraphs)
    else:
        with open(input_file, 'r', encoding='utf-8') as file:
            plaintext = file.read()
    start_time = time.perf_counter()
    ciphertext = encrypt(key, plaintext)

    end_time = time.perf_counter()
    duration = end_time - start_time
    print(f"Encryption time of {input_file}: {duration} seconds")

    with open(output_file, 'w') as file:
        file.write(ciphertext)


def decrypt_file(key, input_file, output_file):
    with open(input_file, 'r') as file:
        ciphertext = file.read()
    start_time = time.perf_counter()
    plaintext = decrypt(key, ciphertext)

    end_time = time.perf_counter()
    duration = end_time - start_time
    print(f"Decryption time of {input_file}: {duration} seconds")

    with open(output_file, 'w') as file:
        file.write(plaintext)


def print_file_content(file_path):
    if file_path.endswith('.docx'):
        doc = Document(file_path)
        content = '\n'.join(paragraph.text for paragraph in doc.paragraphs)
    else:
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
    print(f'Вміст файлу {file_path}:')
    print(content)


def main():
    key_length = 5
    key = generate_random_key(key_length)
    print('Згенерований ключ:', key)

    plaintext = 'This is just the English text!'
    print('Початковий текст:', plaintext)

    ciphertext = encrypt(key, plaintext)
    print('Зашифрований текст:', ciphertext)

    decrypted = decrypt(key, ciphertext)
    print('Розшифрований текст:', decrypted)

    plaintext = 'Це просто український текст!'
    print('Початковий текст:', plaintext)

    ciphertext = encrypt(key, plaintext)
    print('Зашифрований текст:', ciphertext)

    decrypted = decrypt(key, ciphertext)
    print('Розшифрований текст:', decrypted)

    input_file = 'eng.docx'
    output_file = 'eng_encrypted.txt'
    encrypt_file(key, input_file, output_file)
    decrypt_file(key, output_file, 'eng_decrypted.txt')

    #print_file_content(input_file)
    #print_file_content(output_file)
    #print_file_content('eng_decrypted.txt')

    input_file = 'ukr.txt'
    output_file = 'ukr_encrypted.txt'
    encrypt_file(key, input_file, output_file)
    decrypt_file(key, output_file, 'ukr_decrypted.txt')

    #print_file_content(input_file)
    #print_file_content(output_file)
    #print_file_content('ukr_decrypted.txt')


if __name__ == '__main__':
    main()
